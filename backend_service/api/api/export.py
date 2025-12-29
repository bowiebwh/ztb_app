import logging
import os
import re
import uuid
from tempfile import NamedTemporaryFile
from typing import List, Dict, Any
from fastapi import APIRouter, HTTPException
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from minio_client import client, BUCKET, ensure_bucket

router = APIRouter()
logger = logging.getLogger(__name__)
IMAGE_TOKEN_RE = re.compile(r"\[\[IMAGE\|([^\|\]]+)\|([^\]]*)\]\]")


def apply_styles(doc: Document):
    section = doc.sections[0]
    # 页面设置 A4 + 边距
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)

    # 正文字体/段落：仿宋，小四，首行缩进2字符，行距适中（使用常见字体避免回退）
    style = doc.styles["Normal"]
    body_font = "仿宋_GB2312"
    style.font.name = body_font
    for key in ["w:eastAsia", "w:ascii", "w:hAnsi", "w:cs"]:
        style._element.rPr.rFonts.set(qn(key), body_font)
    for key in ["w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"]:
        style._element.rPr.rFonts.attrib.pop(qn(key), None)
    style.font.size = Pt(12)  # 小四
    style.paragraph_format.first_line_indent = Pt(32)
    style.paragraph_format.line_spacing = Pt(26)

    # 标题样式：方正公文小标宋，小二号/更小字号，黑色
    for level, font_name, size_pt, bold, align, space_before, space_after in [
        (1, "方正公文小标宋", 18, True, WD_ALIGN_PARAGRAPH.CENTER, 18, 14),
        (2, "方正公文小标宋", 16, True, WD_ALIGN_PARAGRAPH.LEFT, 16, 10),
        (3, "方正公文小标宋", 14, True, WD_ALIGN_PARAGRAPH.LEFT, 12, 8),
    ]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = font_name
        for key in ["w:eastAsia", "w:ascii", "w:hAnsi", "w:cs"]:
            h._element.rPr.rFonts.set(qn(key), font_name)
        for key in ["w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"]:
            h._element.rPr.rFonts.attrib.pop(qn(key), None)
        h.font.size = Pt(size_pt)
        h.font.bold = bold
        h.font.color.rgb = RGBColor(0, 0, 0)
        h.paragraph_format.space_before = Pt(space_before)
        h.paragraph_format.space_after = Pt(space_after)
        h.paragraph_format.line_spacing = Pt(26)
        h.paragraph_format.alignment = align


def add_cover(
    doc: Document,
    title: str,
    subtitle: str | None = None,
    project_name: str | None = None,
    bidder_name: str | None = None,
    bid_date: str | None = None,
):
    # 大标题
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(120)
    cover.paragraph_format.first_line_indent = Pt(0)
    run = cover.add_run(title)
    run.font.name = "黑体"
    for key in ["w:eastAsia", "w:ascii", "w:hAnsi", "w:cs"]:
        run._element.rPr.rFonts.set(qn(key), "黑体")
    for key in ["w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"]:
        run._element.rPr.rFonts.attrib.pop(qn(key), None)
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    # 副标题（正本）
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.first_line_indent = Pt(0)
        r2 = p2.add_run(subtitle)
        r2.font.name = "黑体"
        for key in ["w:eastAsia", "w:ascii", "w:hAnsi", "w:cs"]:
            r2._element.rPr.rFonts.set(qn(key), "黑体")
        for key in ["w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"]:
            r2._element.rPr.rFonts.attrib.pop(qn(key), None)
        r2.font.size = Pt(16)

    # 间隔：插入空行 + 适度段后距离
    doc.add_paragraph()  # 空行
    doc.add_paragraph()  
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(24)

    # 项目信息
    info_font = "黑体"
    info_size = Pt(16)  # 三号

    def add_info_line(label: str, value: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Pt(16)  # 缩进约1字符
        r_label = p.add_run(label)
        r_label.font.name = info_font
        for key in ["w:eastAsia", "w:ascii", "w:hAnsi", "w:cs"]:
            r_label._element.rPr.rFonts.set(qn(key), info_font)
        for key in ["w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"]:
            r_label._element.rPr.rFonts.attrib.pop(qn(key), None)
        r_label.font.size = info_size
        r_label.font.bold = False

        r_val = p.add_run(value)
        r_val.font.name = info_font
        for key in ["w:eastAsia", "w:ascii", "w:hAnsi", "w:cs"]:
            r_val._element.rPr.rFonts.set(qn(key), info_font)
        for key in ["w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"]:
            r_val._element.rPr.rFonts.attrib.pop(qn(key), None)
        r_val.font.size = info_size

    if project_name:
        add_info_line("项目名称：", project_name)
    # 项目编号占位行
    add_info_line("项目编号：", "")
    if bidder_name:
        add_info_line("投标人名称：", bidder_name)
    if bid_date:
        add_info_line("日      期：", bid_date)

    doc.add_page_break()


def _clean_markdown_line(text: str) -> str:
    t = text.strip()
    if t.startswith("#"):
        t = t.lstrip("#").strip()
    t = t.lstrip("-* ").strip()
    t = t.replace("**", "").replace("__", "")
    return t


def add_sections(doc: Document, sections: List[Dict[str, Any]]):
    def _extract_object_name(url: str | None) -> str | None:
        if not url:
            return None
        if "/download/" in url:
            return url.split("/download/", 1)[-1].lstrip("/")
        return url.lstrip("/")

    def _load_image_tmp(object_name: str | None) -> str | None:
        if not object_name:
            return None
        try:
            resp = client.get_object(BUCKET, object_name)
            _, ext = os.path.splitext(object_name)
            with NamedTemporaryFile(delete=False, suffix=ext or ".img") as tmp:
                data = resp.read()
                tmp.write(data)
                tmp.flush()
                return tmp.name
        except Exception as exc:
            logger.warning("Load image failed | object=%s err=%s", object_name, exc)
            return None

    def _add_paragraph_with_images(text: str):
        # 支持在正文中混排图片标记 [[IMAGE|url|name]]
        parts = IMAGE_TOKEN_RE.split(text or "")
        if not parts:
            return
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for idx, part in enumerate(parts):
            if idx % 2 == 0:
                cleaned = _clean_markdown_line(part)
                if cleaned:
                    p.add_run(cleaned)
            else:
                object_name = _extract_object_name(part)
                tmp_path = _load_image_tmp(object_name)
                if tmp_path:
                    run = p.add_run()
                    try:
                        run.add_picture(tmp_path, width=Cm(14))
                    finally:
                        try:
                            os.unlink(tmp_path)
                        except Exception:
                            pass

    for sec in sections:
        heading = sec.get("heading")
        level = int(sec.get("level", 1))
        body = sec.get("body", "")
        if heading:
            doc.add_heading(_clean_markdown_line(heading), level=level if 1 <= level <= 3 else 1)
        for para in str(body).split("\n"):
            if not para:
                continue
            _add_paragraph_with_images(para)


@router.post("/word")
def export_word(payload: dict):
    """
    接收:
      {
        "title": "投标书标题",
        "subtitle": "可选",
        "sections": [
          { "heading": "1 总体方案", "level": 1, "body": "段落..." },
          ...
        ],
        "content": "当未提供 sections 时的全文文本",
        "filename": "xxx.docx"
      }
    生成 Word 并保存到 MinIO，返回下载 URL。
    """
    title = payload.get("title") or "投标书"
    subtitle = payload.get("subtitle")
    sections = payload.get("sections")
    content = payload.get("content")
    filename = payload.get("filename") or f"{title}.docx"
    if not sections and not content:
        raise HTTPException(status_code=400, detail="content 或 sections 至少提供一个")

    doc = Document()
    apply_styles(doc)
    add_cover(
        doc,
        title,
        subtitle,
        project_name=payload.get("project_name"),
        bidder_name=payload.get("bidder_name"),
        bid_date=payload.get("bid_date"),
    )

    if sections:
        add_sections(doc, sections)
    else:
        doc.add_heading(title, level=1)
        for paragraph in str(content).split("\n"):
            paragraph = _clean_markdown_line(paragraph)
            if paragraph:
                p = doc.add_paragraph(paragraph)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ensure_bucket()
    object_name = f"exports/{uuid.uuid4()}.docx"
    with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        tmp.flush()
        file_size = os.path.getsize(tmp.name)
        with open(tmp.name, "rb") as f:
            client.put_object(
                BUCKET,
                object_name,
                f,
                length=file_size,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    return {
        "url": f"/api/files/download/{object_name}",
        "object_name": object_name,
        "filename": filename,
        "size": file_size,
    }
